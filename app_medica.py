import streamlit as st
from docx import Document
from datetime import datetime

# CONFIGURACIÓN DE PÁGINA
st.set_page_config(page_title="Asistente CEC Van Buren", page_icon="🩺")

st.title("Generador de Documentación CEC - SSVSA")
st.info("Hospital Carlos Van Buren / Servicio de Salud Valparaíso San Antonio")

# FORMULARIO DE DATOS (Lo que el becado rellena)
with st.form("datos_proyecto"):
    st.subheader("Información del Investigador y Proyecto")
    
    col1, col2 = st.columns(2)
    with col1:
        nombre_ip = st.text_input("Nombre Investigador Principal (Becado)")
        run_ip = st.text_input("RUN")
        especialidad = st.text_input("Especialidad / Universidad")
    
    with col2:
        titulo_proyecto = st.text_area("Título exacto del Proyecto")
        unidad_hospital = st.text_input("Unidad o Servicio (ej: Cirugía)")
        duracion = st.text_input("Duración del estudio (ej: 12 meses)")

    st.subheader("Naturaleza del Estudio")
    tipo_estudio = st.radio("Tipo de investigación:", 
                             ["Retrospectivo (Fichas Clínicas)", "Reporte de Caso", "Prospectivo (Pacientes)"])

    submit = st.form_submit_button("Generar Borrador Word")

# LÓGICA DE GENERACIÓN (Doc 1 - Carta Conductora)
# NUEVA LÓGICA MEJORADA PARA EL DOC 1
if submit:
    doc = Document()
    
    # Encabezado oficial
    doc.add_paragraph(f"Valparaíso, {datetime.now().strftime('%d de %B de %Y')}")
    doc.add_paragraph("\nDr. Santiago Parry Ramírez\nPresidente Comité Ético Científico\nServicio de Salud Valparaíso - San Antonio\nPresente")
    
    # Párrafo 1: El qué y el dónde
    p1 = doc.add_paragraph("\nPor medio de la presente carta, solicito a ustedes la evaluación de los aspectos éticos del proyecto de investigación titulado: ")
    p1.add_run(f"“{titulo_proyecto}”").bold = True
    
    doc.add_paragraph(f"Este estudio se llevará a cabo en la unidad de {unidad_hospital} del Hospital Carlos Van Buren. Se presenta como requisito para finalizar la formación en la especialidad de {especialidad}.")

    # Párrafo 2: La justificación técnica (Aquí es donde la IA brillará después)
    doc.add_paragraph(f"La investigación consiste en un análisis de datos con el fin de contribuir al conocimiento científico local y mejorar los estándares de atención en nuestro servicio. Tendrá una duración estimada de {duracion}.")

    # Párrafo 3: Equipo (Dato muy importante para el CEC)
    doc.add_paragraph(f"El estudio será ejecutado junto al equipo de co-investigadores y bajo la supervisión del tutor docente del servicio correspondiente.")

    # Párrafo 4: La "Cláusula de Oro" (Obligatoria)
    doc.add_paragraph("Declaro ante Ud. como Presidente del CEC SSVSA, que a la fecha de presentación no he iniciado ningún tipo de gestión asociada al estudio (contacto con pacientes o acceso a fichas clínicas) sin la autorización previa de este Comité.")

    # Párrafo 5: Listado de lo que se adjunta
    doc.add_paragraph("\nSe adjunta a esta solicitud:")
    listado = [
        "Protocolo de Investigación",
        "Formulario de Valor Social",
        "Carta de aceptación de Jefatura de Servicio",
        "Compromiso del Investigador Principal",
        "Solicitud de Dispensa de Consentimiento (si aplica)"
    ]
    for item in listado:
        doc.add_paragraph(f"• {item}", style='List Bullet')

    doc.add_paragraph("\nAtentamente,\n\n_____________________\n" + f"{nombre_ip}\nRUN: {run_ip}\nInvestigador Principal")

    # Guardar y descargar
    doc_path = "Carta_Conductora_VanBuren_2025.docx"
    doc.save(doc_path)
    
    with open(doc_path, "rb") as file:
        st.download_button("📩 Descargar Carta Profesional Completa", file, file_name=doc_path)
