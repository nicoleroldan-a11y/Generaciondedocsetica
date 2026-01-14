import streamlit as st
from docx import Document
from datetime import datetime
import io

st.set_page_config(page_title="Asistente CEC Pro", page_icon="📝")

st.title("Generador Ético Profesional 🩺")
st.markdown("---")

# FORMULARIO ORGANIZADO
with st.form("main_form"):
    st.subheader("1. Datos del Investigador")
    col1, col2 = st.columns(2)
    with col1:
        nombre = st.text_input("Nombre Completo")
        run = st.text_input("RUT")
    with col2:
        especialidad = st.text_input("Especialidad / Programa")
        hospital = st.selectbox("Hospital", ["Carlos Van Buren", "Gustavo Fricke"])

    st.subheader("2. Detalles del Proyecto")
    titulo = st.text_area("Título del Proyecto")
    resumen = st.text_area("Resumen Breve (La IA lo usará para el Valor Social)")
    
    submit = st.form_submit_button("Generar Set de Documentos")

if submit:
    # --- GENERAR CARTA CONDUCTORA ---
    doc1 = Document()
    doc1.add_paragraph(f"Valparaíso, {datetime.now().strftime('%d de %m de %Y')}")
    doc1.add_paragraph("\nDr. Santiago Parry Ramírez\nPresidente CEC SSVSA\nPresente")
    p = doc1.add_paragraph("\nSolicito evaluación ética del proyecto: ")
    p.add_run(f"“{titulo}”").bold = True
    doc1.add_paragraph(f"\nYo, {nombre}, RUN {run}, en mi calidad de residente de {especialidad}, presento este estudio a realizarse en el Hospital {hospital}...")
    # (Aquí va el resto del texto pro que te pasé antes)
    
    # --- GENERAR DOC 5: VALOR SOCIAL ---
    doc5 = Document()
    doc5.add_heading('DECLARACIÓN SOBRE VALOR SOCIAL', level=1)
    doc5.add_paragraph(f"El trabajo titulado “{titulo}” permitirá conocer la realidad local de nuestra institución.")
    doc5.add_paragraph(f"\nResumen técnico: {resumen}")
    doc5.add_paragraph("\nEl valor social para nosotros es: Contribuir a la mejora de los procesos clínicos y optimización de recursos institucionales basados en evidencia local.")
    
    # DESCARGAS
    st.success("✅ ¡Documentos listos!")
    
    # Preparamos los archivos para descargar sin guardarlos en el disco (en memoria)
    buffer1 = io.BytesIO()
    doc1.save(buffer1)
    st.download_button("📩 Descargar Carta Conductora", buffer1.getvalue(), "Carta_Conductora.docx")
    
    buffer5 = io.BytesIO()
    doc5.save(buffer5)
    st.download_button("📩 Descargar Valor Social (Doc 5)", buffer5.getvalue(), "Valor_Social.docx")
